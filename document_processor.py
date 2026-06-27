"""
document_processor.py — Extract, chunk, and process legal documents
=====================================================================
Supports PDF, DOCX, TXT, and images (OCR).
"""

import re
import json
import uuid
import os
from pathlib import Path
from typing import List, Dict, Optional, Tuple

from config import (
    CHUNK_SIZE, CHUNK_OVERLAP, UPLOAD_DIR,
    ALLOWED_EXTENSIONS, MIN_CHUNK_LENGTH, ENABLE_OCR
)
from utils import logger, clean_text, generate_hash, safe_filename, allowed_file


# ── Text Extraction ───────────────────────────────────────────────────────────

def extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        text_parts = []
        doc = fitz.open(file_path)
        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                text_parts.append(f"[Page {page_num + 1}]\n{text}")
            elif ENABLE_OCR:
                # Try OCR on image-based pages
                ocr_text = _ocr_pdf_page(page)
                if ocr_text:
                    text_parts.append(f"[Page {page_num + 1} (OCR)]\n{ocr_text}")
        doc.close()
        return "\n\n".join(text_parts)
    except ImportError:
        logger.error("PyMuPDF not installed. pip install PyMuPDF")
        return ""
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return ""


def _ocr_pdf_page(page) -> str:
    """OCR a single PDF page using pytesseract."""
    try:
        import fitz
        import pytesseract
        from PIL import Image
        import io

        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img_data = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_data))
        text = pytesseract.image_to_string(img, lang="eng")
        return text.strip()
    except Exception as e:
        logger.warning(f"OCR page error: {e}")
        return ""


def extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX."""
    try:
        from docx import Document
        doc = Document(file_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except ImportError:
        logger.error("python-docx not installed.")
        return ""
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""


def extract_from_txt(file_path: str) -> str:
    """Extract text from TXT file."""
    for encoding in ["utf-8", "utf-16", "latin-1", "ascii"]:
        try:
            with open(file_path, "r", encoding=encoding, errors="replace") as f:
                return f.read()
        except Exception:
            continue
    return ""


def extract_from_image(file_path: str) -> str:
    """Extract text from image using pytesseract + EasyOCR."""
    text = ""

    # Try pytesseract
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img, lang="eng")
        if text.strip():
            return text.strip()
    except Exception as e:
        logger.warning(f"pytesseract failed: {e}")

    # Fallback to EasyOCR (optional heavy dependency)
    try:
        import importlib
        easyocr = importlib.import_module("easyocr")  # type: ignore[assignment]
        reader = easyocr.Reader(["en", "hi"], gpu=False)
        results = reader.readtext(file_path)
        text = " ".join([r[1] for r in results])
    except Exception as e:
        logger.warning(f"EasyOCR failed (install easyocr for better OCR): {e}")

    return text.strip()


def extract_text(file_path: str, file_type: str = None) -> str:
    """Auto-detect file type and extract text."""
    path = Path(file_path)
    ext = file_type or path.suffix.lower().lstrip(".")

    extractors = {
        "pdf": extract_from_pdf,
        "docx": extract_from_docx,
        "doc": extract_from_docx,
        "txt": extract_from_txt,
        "png": extract_from_image,
        "jpg": extract_from_image,
        "jpeg": extract_from_image,
        "tiff": extract_from_image,
        "tif": extract_from_image,
    }

    extractor = extractors.get(ext)
    if not extractor:
        logger.warning(f"Unsupported file type: {ext}")
        return ""

    text = extractor(str(file_path))
    return clean_text(text)


# ── Text Chunking ─────────────────────────────────────────────────────────────

def chunk_text(text: str, chunk_size: int = CHUNK_SIZE,
               overlap: int = CHUNK_OVERLAP) -> List[str]:
    """
    Split text into overlapping chunks at sentence boundaries.
    """
    if not text or len(text) < MIN_CHUNK_LENGTH:
        return [text] if text else []

    # Split into sentences first
    sentences = re.split(r'(?<=[.!?])\s+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) >= 10]

    chunks = []
    current_chunk = []
    current_len = 0

    for sentence in sentences:
        sent_len = len(sentence.split())

        if current_len + sent_len > chunk_size and current_chunk:
            chunk_text_str = " ".join(current_chunk)
            if len(chunk_text_str) >= MIN_CHUNK_LENGTH:
                chunks.append(chunk_text_str)

            # Overlap: keep last N words
            overlap_words = " ".join(current_chunk).split()[-overlap:]
            current_chunk = overlap_words
            current_len = len(overlap_words)

        current_chunk.append(sentence)
        current_len += sent_len

    # Last chunk
    if current_chunk:
        chunk_text_str = " ".join(current_chunk)
        if len(chunk_text_str) >= MIN_CHUNK_LENGTH:
            chunks.append(chunk_text_str)

    return chunks if chunks else [text[:chunk_size * 5]]


def chunk_by_section(text: str) -> List[Dict]:
    """
    Chunk document by legal sections/chapters.
    Returns chunks with section metadata.
    """
    section_pattern = re.compile(
        r'(?:SECTION|Section|CHAPTER|Chapter|PART|Part)\s+\d+[A-Z]?[.:]?\s*[A-Z][^\n]{0,100}',
        re.MULTILINE
    )

    splits = list(section_pattern.finditer(text))
    if len(splits) < 2:
        # Fall back to regular chunking
        return [{"section": "Main", "text": c} for c in chunk_text(text)]

    chunks = []
    for i, match in enumerate(splits):
        section_header = match.group(0).strip()
        start = match.end()
        end = splits[i + 1].start() if i + 1 < len(splits) else len(text)
        section_text = text[start:end].strip()

        # Further chunk large sections
        sub_chunks = chunk_text(section_text)
        for sub in sub_chunks:
            chunks.append({"section": section_header, "text": sub})

    return chunks


# ── File Handling ─────────────────────────────────────────────────────────────

def save_uploaded_file(file_obj, filename: str) -> Tuple[str, str]:
    """
    Save uploaded file to upload directory.
    Returns (saved_path, unique_filename).
    """
    safe_name = safe_filename(filename)
    unique_name = f"{uuid.uuid4().hex[:8]}_{safe_name}"
    save_path = Path(UPLOAD_DIR) / unique_name
    file_obj.save(str(save_path))
    logger.info(f"File saved: {save_path}")
    return str(save_path), unique_name


def process_document(file_path: str, filename: str, doc_id: int = None) -> Dict:
    """
    Full document processing pipeline:
    1. Extract text
    2. Chunk text
    3. Extract metadata
    Returns dict with all extracted data.
    """
    path = Path(file_path)
    file_type = path.suffix.lower().lstrip(".")
    file_size = path.stat().st_size if path.exists() else 0

    # Extract text
    raw_text = extract_text(file_path, file_type)

    if not raw_text:
        return {
            "success": False,
            "error": "Could not extract text from document.",
            "file_path": file_path,
        }

    # Chunk for vector storage
    chunks = chunk_text(raw_text)

    # Hash for duplicate detection
    content_hash = generate_hash(raw_text)

    result = {
        "success": True,
        "file_path": file_path,
        "filename": filename,
        "file_type": file_type,
        "file_size": file_size,
        "raw_text": raw_text,
        "chunks": chunks,
        "chunk_count": len(chunks),
        "word_count": len(raw_text.split()),
        "char_count": len(raw_text),
        "content_hash": content_hash,
        "ocr_processed": file_type in {"png", "jpg", "jpeg", "tiff"},
    }

    return result


def get_document_preview(file_path: str, max_chars: int = 1000) -> str:
    """Return a preview of document text."""
    text = extract_text(file_path)
    return text[:max_chars] + ("..." if len(text) > max_chars else "")
